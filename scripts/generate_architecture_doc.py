from pathlib import Path
from datetime import date
import textwrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "architecture_outputs"
OUT_DIR.mkdir(exist_ok=True)
DOCX = ROOT / "沪面冲鸭技术架构设计文档.docx"

W, H = 1600, 1000
BG = (248, 249, 251)
INK = (18, 24, 38)
MUTED = (102, 112, 133)
LINE = (218, 225, 235)
CARD = (255, 255, 255)
ORANGE = (255, 122, 24)
BLUE = (42, 101, 214)
GREEN = (22, 163, 74)
PURPLE = (126, 87, 194)
RED = (224, 72, 72)


def font(size, bold=False):
    paths = [
        "C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf" if bold else "C:/Windows/Fonts/simsun.ttc",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
    ]
    for path in paths:
        if Path(path).exists():
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


F14, F16, F18, F20, F24, F30 = [font(s) for s in [14, 16, 18, 20, 24, 30]]
FB14, FB16, FB18, FB20, FB24, FB30, FB36 = [font(s, True) for s in [14, 16, 18, 20, 24, 30, 36]]


def rr(d, box, r=24, fill=CARD, outline=LINE, width=2):
    d.rounded_rectangle(box, r, fill=fill, outline=outline, width=width)


def text(d, xy, s, f=F18, fill=INK, anchor=None, align="left"):
    d.text(xy, s, font=f, fill=fill, anchor=anchor, align=align)


def wrap_lines(d, s, max_width, f):
    lines, buf = [], ""
    for ch in s:
        test = buf + ch
        if d.textlength(test, font=f) <= max_width:
            buf = test
        else:
            if buf:
                lines.append(buf)
            buf = ch
    if buf:
        lines.append(buf)
    return lines


def box(d, xywh, title, body=None, color=BLUE, icon=None):
    x, y, w, h = xywh
    rr(d, (x, y, x + w, y + h), 22, CARD, LINE, 2)
    d.rounded_rectangle((x, y, x + w, y + 12), 8, fill=color, outline=color)
    if icon:
        text(d, (x + 28, y + 42), icon, FB20, color, "mm")
        tx = x + 52
    else:
        tx = x + 24
    text(d, (tx, y + 35), title, FB20, INK, "lm")
    if body:
        yy = y + 70
        for line in body if isinstance(body, list) else wrap_lines(d, body, w - 44, F16):
            text(d, (x + 24, yy), line, F16, MUTED)
            yy += 28


def arrow(d, start, end, color=(93, 103, 118), width=3):
    d.line((start[0], start[1], end[0], end[1]), fill=color, width=width)
    angle = __import__("math").atan2(end[1] - start[1], end[0] - start[0])
    import math
    size = 12
    p1 = (end[0] - size * math.cos(angle - math.pi / 6), end[1] - size * math.sin(angle - math.pi / 6))
    p2 = (end[0] - size * math.cos(angle + math.pi / 6), end[1] - size * math.sin(angle + math.pi / 6))
    d.polygon([end, p1, p2], fill=color)


def title(d, s, subtitle=None):
    text(d, (80, 58), s, FB36, INK)
    if subtitle:
        text(d, (80, 105), subtitle, F18, MUTED)


def save(img, name):
    path = OUT_DIR / name
    img.convert("RGB").save(path, quality=95)
    return path


def draw_overall_architecture():
    img = Image.new("RGBA", (W, H), BG + (255,))
    d = ImageDraw.Draw(img)
    title(d, "总体架构图", "浏览器前端、后端 API、数据库、文件存储与 DeepSeek AI 的整体关系")
    box(d, (80, 190, 310, 180), "用户浏览器", ["React + Vite + TS", "录音采集 MediaRecorder", "沉浸式面试 UI"], ORANGE)
    box(d, (520, 160, 380, 240), "前端 Web App", ["页面路由与组件", "Zustand 状态管理", "题库/组题/面试流程", "调用后端 API"], BLUE)
    box(d, (1030, 145, 420, 270), "后端 API 服务", ["NestJS / Node.js", "Auth / Question / Interview", "Upload / AI / Report", "统一鉴权与日志"], GREEN)
    box(d, (160, 560, 320, 210), "对象存储", ["开发期：本地 uploads", "正式期：阿里云 OSS", "保存 WebM/Opus 录音"], PURPLE)
    box(d, (640, 560, 320, 210), "PostgreSQL", ["用户、题库、岗位信息", "面试记录、作答、AI结果", "Prisma ORM"], BLUE)
    box(d, (1120, 560, 320, 210), "DeepSeek API", ["AI 单题点评", "AI 答题思路", "AI 生题", "整场报告"], RED)
    arrow(d, (390, 280), (520, 280))
    arrow(d, (900, 280), (1030, 280))
    arrow(d, (1240, 415), (1280, 560))
    arrow(d, (1130, 415), (800, 560))
    arrow(d, (1080, 415), (320, 560))
    arrow(d, (640, 560), (390, 370))
    text(d, (430, 255), "HTTPS", FB16, MUTED)
    text(d, (936, 255), "REST API", FB16, MUTED)
    return save(img, "01_overall_architecture.png")


def draw_frontend_backend_modules():
    img = Image.new("RGBA", (W, H), BG + (255,))
    d = ImageDraw.Draw(img)
    title(d, "前后端模块划分图", "按 PRD 页面、业务能力和后端服务边界拆分模块")
    box(d, (80, 170, 620, 680), "前端模块", [
        "App Router：登录、首页、题库、岗位、面试、复盘、个人中心",
        "UI Components：导航、三栏题库、自由组题、面试控件、报告面板",
        "Interview Store：模式、题号、倒计时、录音、题目显隐",
        "Recorder Service：MediaRecorder 分段录音、本地预览",
        "API Client：统一请求、错误提示、登录态处理",
    ], BLUE)
    box(d, (900, 170, 620, 680), "后端模块", [
        "AuthModule：邮箱注册/登录、JWT、用户身份",
        "QuestionModule：真题、套题、分类树、用户自建题",
        "JobProfileModule：岗位信息保存与读取",
        "InterviewModule：创建面试、结束面试、历史记录",
        "UploadModule：录音上传、文件元数据",
        "AIModule：DeepSeek 调用、Prompt 管理、失败兜底",
        "ReportModule：复盘页数据聚合、报告导出预留",
        "MembershipModule：会员/题库权限预留",
    ], GREEN)
    arrow(d, (700, 330), (900, 330))
    arrow(d, (900, 390), (700, 390))
    text(d, (800, 312), "API 请求", FB16, MUTED, "mm")
    text(d, (800, 430), "JSON 响应", FB16, MUTED, "mm")
    return save(img, "02_module_architecture.png")


def draw_activity_interview():
    img = Image.new("RGBA", (W, H), BG + (255,))
    d = ImageDraw.Draw(img)
    title(d, "UML 活动图：模拟面试流程", "覆盖听题模式、看题模式、录音分段和复盘生成")
    steps = [
        ("登录/进入首页", 120, 160),
        ("选择套题、时间、模式", 120, 300),
        ("创建面试 Session", 120, 440),
        ("判断模式", 580, 440),
        ("听题：播放题目语音", 580, 260),
        ("看题：显示题目阅读层", 580, 620),
        ("开始/持续录音转写", 1020, 440),
        ("点击下一题或总时间结束", 1020, 590),
        ("保存本题录音与文本", 1020, 740),
        ("生成 AI 复盘", 580, 800),
    ]
    for label, x, y in steps:
        box(d, (x, y, 280, 78), label, None, ORANGE if "判断" in label else BLUE)
    arrow(d, (260, 238), (260, 300))
    arrow(d, (260, 378), (260, 440))
    arrow(d, (400, 479), (580, 479))
    arrow(d, (720, 440), (720, 338))
    arrow(d, (720, 518), (720, 620))
    arrow(d, (860, 299), (1020, 459))
    arrow(d, (860, 659), (1020, 479))
    arrow(d, (1160, 518), (1160, 590))
    arrow(d, (1160, 668), (1160, 740))
    arrow(d, (1020, 779), (860, 839))
    text(d, (820, 410), "听题", FB16, MUTED)
    text(d, (820, 568), "看题", FB16, MUTED)
    return save(img, "03_activity_interview.png")


def draw_sequence_ai_review():
    img = Image.new("RGBA", (W, H), BG + (255,))
    d = ImageDraw.Draw(img)
    title(d, "UML 时序图：提交作答与 AI 点评", "前端结束面试后，后端保存数据并调用 DeepSeek 生成复盘")
    actors = [("前端", 180), ("后端 API", 520), ("PostgreSQL", 860), ("对象存储", 1160), ("DeepSeek", 1400)]
    for name, x in actors:
        rr(d, (x - 80, 160, x + 80, 215), 14, CARD, LINE)
        text(d, (x, 188), name, FB16, INK, "mm")
        d.line((x, 215, x, 880), fill=(200, 205, 214), width=2)
    events = [
        (250, 180, 520, "POST /interviews/:id/answers"),
        (330, 520, 1160, "上传录音文件"),
        (410, 520, 860, "保存 answer/audio 元数据"),
        (500, 520, 1400, "调用 DeepSeek 生成评语/思路"),
        (590, 1400, 520, "返回 AI JSON"),
        (670, 520, 860, "保存 ai_reviews"),
        (760, 520, 180, "返回复盘数据"),
    ]
    for y, x1, x2, label in events:
        arrow(d, (x1, y), (x2, y), BLUE if x2 > x1 else GREEN)
        text(d, ((x1 + x2) / 2, y - 28), label, F14, MUTED, "mm")
    return save(img, "04_sequence_ai_review.png")


def draw_er_diagram():
    img = Image.new("RGBA", (W, H), BG + (255,))
    d = ImageDraw.Draw(img)
    title(d, "数据库 ER 关系图", "PostgreSQL + Prisma 的核心数据对象关系")
    tables = {
        "users": (80, 170, ["id PK", "username", "email", "created_at"]),
        "job_profiles": (430, 170, ["id PK", "user_id FK", "job_title", "unit_name"]),
        "questions": (780, 170, ["id PK", "title", "content", "type", "source"]),
        "question_sets": (1130, 170, ["id PK", "title", "region", "exam_date"]),
        "custom_questions": (80, 560, ["id PK", "user_id FK", "content", "type"]),
        "interviews": (430, 560, ["id PK", "user_id FK", "mode", "status"]),
        "interview_answers": (780, 560, ["id PK", "interview_id FK", "question_id FK", "transcript"]),
        "audio_files": (1130, 560, ["id PK", "answer_id FK", "file_url", "duration"]),
        "ai_reviews": (780, 800, ["id PK", "answer_id FK", "comment", "score"]),
    }
    centers = {}
    for name, (x, y, cols) in tables.items():
        rr(d, (x, y, x + 290, y + 220), 18, CARD, LINE)
        d.rectangle((x, y, x + 290, y + 48), fill=(239, 246, 255), outline=LINE)
        text(d, (x + 20, y + 25), name, FB18, INK, "lm")
        yy = y + 70
        for col in cols:
            text(d, (x + 22, yy), col, F15 if False else F16, MUTED)
            yy += 30
        centers[name] = (x + 145, y + 110)
    rels = [
        ("users", "job_profiles"),
        ("users", "custom_questions"),
        ("users", "interviews"),
        ("questions", "interview_answers"),
        ("question_sets", "questions"),
        ("interviews", "interview_answers"),
        ("interview_answers", "audio_files"),
        ("interview_answers", "ai_reviews"),
    ]
    for a, b in rels:
        arrow(d, centers[a], centers[b], (93, 103, 118), 2)
    return save(img, "05_er_diagram.png")


def draw_deployment():
    img = Image.new("RGBA", (W, H), BG + (255,))
    d = ImageDraw.Draw(img)
    title(d, "部署架构图", "MVP 与正式版均可落地的部署拓扑")
    box(d, (90, 180, 340, 170), "浏览器", ["Chrome / Edge 桌面端", "暂不支持移动端模拟"], ORANGE)
    box(d, (520, 180, 360, 170), "静态前端", ["Nginx / OSS 静态站点", "React build 产物"], BLUE)
    box(d, (1080, 180, 360, 170), "API 服务", ["NestJS Docker 服务", "环境变量管理密钥"], GREEN)
    box(d, (260, 570, 300, 170), "PostgreSQL", ["RDS / 云数据库", "Prisma migration"], BLUE)
    box(d, (650, 570, 300, 170), "阿里云 OSS", ["录音文件", "报告文件预留"], PURPLE)
    box(d, (1040, 570, 300, 170), "DeepSeek API", ["公网 API", "后端代理调用"], RED)
    arrow(d, (430, 265), (520, 265))
    arrow(d, (880, 265), (1080, 265))
    arrow(d, (1190, 350), (410, 570))
    arrow(d, (1220, 350), (800, 570))
    arrow(d, (1260, 350), (1190, 570))
    text(d, (475, 240), "HTTPS", FB16, MUTED)
    text(d, (970, 240), "API", FB16, MUTED)
    return save(img, "06_deployment_architecture.png")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, value, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(value))
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.size = Pt(8.5)
    r.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, True)
        set_cell_shading(table.rows[0].cells[i], "F2F4F7")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val)
    if widths:
        for row in table.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def setup_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    styles = doc.styles
    styles["Normal"].font.name = "微软雅黑"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    styles["Normal"].font.size = Pt(10.5)


def heading(doc, s, level=1):
    p = doc.add_heading(s, level=level)
    for run in p.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.color.rgb = RGBColor(17, 24, 39)
    return p


def para(doc, s=""):
    p = doc.add_paragraph(s)
    for run in p.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.size = Pt(10.5)
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def add_image(doc, path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(6.9))
    c = doc.add_paragraph(caption)
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in c.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(102, 112, 133)


def cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("沪面冲鸭技术架构设计文档")
    r.bold = True
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.size = Pt(24)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("上海事业单位面试模拟网站｜前端、后端、数据库、AI、部署方案")
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(102, 112, 133)
    for line in ["版本：V1.0", f"日期：{date.today().isoformat()}", "技术栈：React + Vite + TypeScript + Zustand / NestJS + Prisma / PostgreSQL / OSS / DeepSeek API"]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()


def build():
    diagrams = [
        draw_overall_architecture(),
        draw_frontend_backend_modules(),
        draw_activity_interview(),
        draw_sequence_ai_review(),
        draw_er_diagram(),
        draw_deployment(),
    ]
    doc = Document()
    setup_doc(doc)
    cover(doc)

    heading(doc, "1. 架构目标", 1)
    para(doc, "本文档基于《上海事业单位面试模拟网站 PRD_完善版》编写，用于明确系统从 MVP 到 V1.0 的技术架构、模块边界、UML 建模、数据库设计和部署方案。")
    bullets(doc, [
        "支持桌面端优先的沉浸式模拟面试体验。",
        "支持题库、自由组题、岗位信息、录音转写、AI 复盘和历史记录。",
        "MVP 可快速落地，正式版可平滑扩展后端、对象存储、会员和管理后台。",
        "DeepSeek API 由后端统一代理调用，避免 API Key 暴露在浏览器。",
    ])

    heading(doc, "2. 技术栈选型", 1)
    add_table(doc, ["层级", "技术选型", "说明"], [
        ["前端", "React + Vite + TypeScript", "适合复杂交互与组件化页面，开发速度快，类型安全。"],
        ["状态管理", "Zustand", "管理面试流程、录音状态、题目显隐、用户登录态等全局状态。"],
        ["后端", "NestJS + Node.js", "模块化清晰，适合 Auth、题库、面试、AI、上传、会员等能力扩展。"],
        ["ORM", "Prisma", "类型安全、迁移清晰，适合 PostgreSQL 数据建模。"],
        ["数据库", "PostgreSQL", "关系型数据稳定可靠，适合用户、题库、面试记录、AI 结果。"],
        ["文件存储", "开发期本地 uploads；正式期阿里云 OSS", "录音文件不入库，数据库只保存 URL 和元数据。"],
        ["AI", "DeepSeek API", "用于单题点评、答题思路、岗位 AI 生题、整场报告。"],
        ["语音", "MVP Web Speech API + MediaRecorder；正式期服务端 ASR", "MVP 快速跑通，后续可接阿里云/讯飞/腾讯云 ASR。"],
        ["部署", "前端静态站点 + NestJS API + PostgreSQL + OSS", "国内访问优先可选阿里云 ECS/RDS/OSS。"],
    ], widths=[1.1, 2.2, 3.3])

    heading(doc, "3. 总体架构", 1)
    para(doc, "系统采用前后端分离架构。浏览器负责页面展示、面试流程控制和录音采集；后端负责鉴权、业务数据、文件上传、AI 调用和报告聚合；数据库保存结构化数据；对象存储保存录音文件；DeepSeek API 提供 AI 能力。")
    add_image(doc, diagrams[0], "图 1 总体架构图")

    heading(doc, "4. 前后端模块设计", 1)
    add_image(doc, diagrams[1], "图 2 前后端模块划分图")
    heading(doc, "4.1 前端页面路由", 2)
    add_table(doc, ["路由", "页面", "说明"], [
        ["/login", "登录注册页", "用户名 + 邮箱登录，后续可扩展密码/验证码。"],
        ["/", "开始面试首页", "品牌入口，进入面试配置。"],
        ["/interview/setup", "面试配置页", "选择套题、时间、听题/看题模式。"],
        ["/interview/session/:id", "模拟面试页", "沉浸式背景、录音、倒计时、下一题。"],
        ["/interview/review/:id", "面试复盘页", "查看录音、转写、AI 评语和答题思路。"],
        ["/question-bank", "题库页", "分类树、题目列表、自由组题。"],
        ["/job-profile", "岗位信息页", "岗位信息保存、AI 生题、加入题库。"],
        ["/profile", "个人中心", "历史面试记录。"],
    ], widths=[1.8, 1.7, 3.0])

    heading(doc, "4.2 后端模块", 2)
    add_table(doc, ["模块", "职责"], [
        ["AuthModule", "注册、登录、JWT 鉴权、当前用户信息。"],
        ["QuestionModule", "题库分类树、真题、套题、自建题、题目查询。"],
        ["JobProfileModule", "岗位信息保存、读取、更新。"],
        ["InterviewModule", "创建面试、结束面试、历史记录、复盘数据聚合。"],
        ["AnswerModule", "每题作答记录、题目快照、转写文本、时长。"],
        ["UploadModule", "录音文件上传、文件元数据、对象存储适配。"],
        ["AIModule", "DeepSeek 调用、Prompt 管理、AI 结果解析与兜底。"],
        ["ReportModule", "整场报告生成、导出预留。"],
        ["MembershipModule", "会员与题库权限预留，MVP 可不启用。"],
    ], widths=[1.8, 4.7])

    heading(doc, "5. UML 建模", 1)
    heading(doc, "5.1 活动图：模拟面试流程", 2)
    add_image(doc, diagrams[2], "图 3 UML 活动图 - 模拟面试流程")
    heading(doc, "5.2 时序图：提交作答与 AI 点评", 2)
    add_image(doc, diagrams[3], "图 4 UML 时序图 - 提交作答与 AI 点评")

    heading(doc, "5.3 用例模型", 2)
    add_table(doc, ["参与者", "用例"], [
        ["未登录用户", "查看首页、登录/注册。"],
        ["已登录考生", "浏览题库、自由组题、填写岗位信息、开始模拟、录音作答、查看 AI 复盘、查看历史记录。"],
        ["AI 服务", "生成单题点评、答题思路、岗位题、整场报告。"],
        ["系统管理员（后续）", "维护题库、审核真题、管理用户、配置会员权限。"],
    ], widths=[1.7, 4.8])

    heading(doc, "5.4 领域模型", 2)
    add_table(doc, ["领域对象", "说明", "关键关系"], [
        ["User", "注册用户", "拥有 JobProfile、CustomQuestion、Interview。"],
        ["Question", "标准题库题目", "可属于 QuestionSet，可被 InterviewAnswer 引用。"],
        ["QuestionSet", "套题", "包含多道 Question。"],
        ["Interview", "一次模拟面试", "属于 User，包含多条 InterviewAnswer。"],
        ["InterviewAnswer", "某题作答记录", "关联 Interview、Question、AudioFile、AIReview。"],
        ["AudioFile", "录音文件元数据", "关联 InterviewAnswer。"],
        ["AIReview", "AI 分析结果", "关联 InterviewAnswer，可扩展整场报告。"],
    ], widths=[1.5, 2.2, 2.8])

    heading(doc, "6. 数据库设计", 1)
    add_image(doc, diagrams[4], "图 5 数据库 ER 关系图")
    heading(doc, "6.1 数据表清单", 2)
    add_table(doc, ["表名", "用途"], [
        ["users", "用户基础信息与登录身份。"],
        ["job_profiles", "用户岗位信息。"],
        ["questions", "标准题库题目。"],
        ["question_sets", "套题信息。"],
        ["question_set_items", "套题与题目的多对多顺序关系。"],
        ["custom_questions", "用户手动添加或 AI 生成后加入的专属题目。"],
        ["interviews", "一次面试模拟记录。"],
        ["interview_answers", "每道题的作答记录与题目快照。"],
        ["audio_files", "录音文件元数据。"],
        ["ai_reviews", "逐题 AI 点评和答题思路。"],
        ["ai_reports", "整场 AI 报告，V1.0 可启用。"],
        ["memberships", "会员权限预留。"],
    ], widths=[2.0, 4.5])

    heading(doc, "6.2 核心数据表字段", 2)
    table_defs = {
        "users": [
            ("id", "uuid", "PK", "用户 ID"),
            ("username", "varchar(64)", "not null", "用户名"),
            ("email", "varchar(255)", "unique, not null", "邮箱"),
            ("password_hash", "varchar(255)", "nullable", "后续扩展密码登录"),
            ("created_at", "timestamp", "not null", "创建时间"),
            ("updated_at", "timestamp", "not null", "更新时间"),
        ],
        "job_profiles": [
            ("id", "uuid", "PK", "岗位信息 ID"),
            ("user_id", "uuid", "FK users.id", "所属用户"),
            ("job_title", "varchar(128)", "not null", "岗位名称"),
            ("unit_name", "varchar(128)", "not null", "岗位所在单位"),
            ("requirements", "text", "not null", "岗位要求"),
            ("extra_info", "text", "nullable", "岗位其他信息"),
            ("created_at", "timestamp", "not null", "创建时间"),
            ("updated_at", "timestamp", "not null", "更新时间"),
        ],
        "questions": [
            ("id", "uuid", "PK", "题目 ID"),
            ("title", "varchar(255)", "not null", "题目标题"),
            ("content", "text", "not null", "题目正文"),
            ("type", "varchar(64)", "not null", "综合分析/人际关系/应急应变/组织计划/岗位匹配/专业题"),
            ("source", "varchar(64)", "not null", "真题/AI生成/用户自建"),
            ("region", "varchar(64)", "nullable", "地区"),
            ("unit_name", "varchar(128)", "nullable", "单位"),
            ("exam_date", "date", "nullable", "考试日期"),
            ("rules", "text", "nullable", "考场规则"),
            ("created_at", "timestamp", "not null", "创建时间"),
        ],
        "question_sets": [
            ("id", "uuid", "PK", "套题 ID"),
            ("title", "varchar(255)", "not null", "套题标题"),
            ("source", "varchar(64)", "not null", "真题/自建/AI"),
            ("region", "varchar(64)", "nullable", "地区"),
            ("unit_name", "varchar(128)", "nullable", "单位"),
            ("exam_date", "date", "nullable", "考试日期"),
            ("rules", "text", "nullable", "考场规则"),
            ("created_at", "timestamp", "not null", "创建时间"),
        ],
        "question_set_items": [
            ("id", "uuid", "PK", "关系 ID"),
            ("question_set_id", "uuid", "FK question_sets.id", "套题 ID"),
            ("question_id", "uuid", "FK questions.id", "题目 ID"),
            ("sort_order", "int", "not null", "题目顺序"),
        ],
        "custom_questions": [
            ("id", "uuid", "PK", "自建题 ID"),
            ("user_id", "uuid", "FK users.id", "所属用户"),
            ("content", "text", "not null", "题目正文"),
            ("type", "varchar(64)", "nullable", "题型"),
            ("source", "varchar(64)", "not null", "manual/ai"),
            ("created_at", "timestamp", "not null", "创建时间"),
        ],
        "interviews": [
            ("id", "uuid", "PK", "面试 ID"),
            ("user_id", "uuid", "FK users.id", "所属用户"),
            ("mode", "varchar(32)", "not null", "listen/read"),
            ("source_type", "varchar(64)", "not null", "question_set/free/custom/ai"),
            ("total_seconds", "int", "not null", "总倒计时秒数"),
            ("status", "varchar(32)", "not null", "created/running/finished/canceled"),
            ("started_at", "timestamp", "nullable", "开始时间"),
            ("ended_at", "timestamp", "nullable", "结束时间"),
            ("created_at", "timestamp", "not null", "创建时间"),
        ],
        "interview_answers": [
            ("id", "uuid", "PK", "作答 ID"),
            ("interview_id", "uuid", "FK interviews.id", "面试 ID"),
            ("question_id", "uuid", "nullable FK questions.id", "题库题 ID"),
            ("question_content_snapshot", "text", "not null", "作答时题目快照"),
            ("transcript", "text", "nullable", "转写文本"),
            ("duration_seconds", "int", "nullable", "录音时长"),
            ("sort_order", "int", "not null", "题目顺序"),
            ("created_at", "timestamp", "not null", "创建时间"),
        ],
        "audio_files": [
            ("id", "uuid", "PK", "音频 ID"),
            ("user_id", "uuid", "FK users.id", "所属用户"),
            ("interview_id", "uuid", "FK interviews.id", "面试 ID"),
            ("answer_id", "uuid", "FK interview_answers.id", "作答 ID"),
            ("file_url", "text", "not null", "录音地址"),
            ("mime_type", "varchar(64)", "not null", "audio/webm"),
            ("size_bytes", "bigint", "nullable", "文件大小"),
            ("duration_seconds", "int", "nullable", "时长"),
            ("created_at", "timestamp", "not null", "创建时间"),
        ],
        "ai_reviews": [
            ("id", "uuid", "PK", "AI 点评 ID"),
            ("interview_answer_id", "uuid", "FK interview_answers.id", "作答 ID"),
            ("score", "int", "nullable", "单题分数"),
            ("comment", "text", "not null", "AI 评语"),
            ("thinking", "text", "not null", "AI 答题思路"),
            ("strengths", "jsonb", "nullable", "优点列表"),
            ("weaknesses", "jsonb", "nullable", "问题列表"),
            ("suggestions", "jsonb", "nullable", "建议列表"),
            ("created_at", "timestamp", "not null", "创建时间"),
        ],
        "ai_reports": [
            ("id", "uuid", "PK", "报告 ID"),
            ("interview_id", "uuid", "FK interviews.id", "面试 ID"),
            ("total_score", "int", "nullable", "总分"),
            ("match_score", "int", "nullable", "岗位匹配度"),
            ("stability_score", "int", "nullable", "表达稳定度"),
            ("summary", "text", "not null", "总评"),
            ("details", "jsonb", "nullable", "维度详情"),
            ("created_at", "timestamp", "not null", "创建时间"),
        ],
        "memberships": [
            ("id", "uuid", "PK", "会员 ID"),
            ("user_id", "uuid", "FK users.id", "用户 ID"),
            ("plan", "varchar(64)", "not null", "free/pro"),
            ("started_at", "timestamp", "nullable", "开始时间"),
            ("expired_at", "timestamp", "nullable", "过期时间"),
            ("status", "varchar(32)", "not null", "active/expired/canceled"),
        ],
    }
    for table, rows in table_defs.items():
        heading(doc, table, 3)
        add_table(doc, ["字段", "类型", "约束", "说明"], rows, widths=[1.5, 1.5, 1.7, 2.0])

    heading(doc, "6.3 关键索引与约束", 2)
    bullets(doc, [
        "users.email 建唯一索引，避免重复注册。",
        "job_profiles.user_id 建唯一索引，MVP 默认每个用户一份当前岗位信息。",
        "questions.type、questions.region、questions.exam_date 建普通索引，支持题库筛选。",
        "question_set_items(question_set_id, sort_order) 建联合索引，保证套题题目顺序。",
        "interviews(user_id, created_at) 建索引，支持个人中心按时间倒序查询。",
        "interview_answers(interview_id, sort_order) 建联合索引，支持复盘页按题号展示。",
        "custom_questions(user_id) 需限制单用户最多 10 道，可在业务层校验。",
    ])

    heading(doc, "7. API 设计", 1)
    add_table(doc, ["方法", "路径", "说明"], [
        ["POST", "/api/auth/register", "注册用户。"],
        ["POST", "/api/auth/login", "登录并返回 token。"],
        ["GET", "/api/auth/me", "获取当前用户。"],
        ["GET", "/api/questions/tree", "获取题库分类树。"],
        ["GET", "/api/questions", "按分类/地区/题型查询题目。"],
        ["POST", "/api/questions/custom", "新增我的专属题目。"],
        ["PUT", "/api/questions/custom/:id", "编辑我的专属题目。"],
        ["DELETE", "/api/questions/custom/:id", "删除我的专属题目。"],
        ["GET", "/api/job-profile", "获取岗位信息。"],
        ["POST", "/api/job-profile", "保存岗位信息。"],
        ["POST", "/api/interviews", "创建面试。"],
        ["GET", "/api/interviews", "查询历史面试。"],
        ["GET", "/api/interviews/:id", "获取复盘详情。"],
        ["POST", "/api/interviews/:id/answers", "保存某题作答。"],
        ["POST", "/api/interviews/:id/finish", "结束面试并触发报告生成。"],
        ["POST", "/api/uploads/audio", "上传录音文件。"],
        ["POST", "/api/ai/review-answer", "生成单题点评。"],
        ["POST", "/api/ai/generate-questions", "根据岗位信息生成 10 道题。"],
        ["POST", "/api/ai/generate-report", "生成整场报告。"],
    ], widths=[0.8, 2.4, 3.3])

    heading(doc, "8. AI 与录音方案", 1)
    bullets(doc, [
        "DeepSeek API 只能由后端调用，前端不保存或暴露 API Key。",
        "Prompt 分为单题点评、答题思路、岗位 AI 生题、整场报告四类，统一由 AIModule 管理。",
        "MVP 使用 MediaRecorder 保存 WebM/Opus 音频，Web Speech API 做实时转写。",
        "正式版可将录音上传后调用服务端 ASR，提升识别稳定性和浏览器兼容性。",
        "AI 失败时不影响录音、转写和历史记录保存，复盘页显示可重试状态。",
    ])

    heading(doc, "9. 部署架构", 1)
    add_image(doc, diagrams[5], "图 6 部署架构图")
    add_table(doc, ["环境", "前端", "后端", "数据库", "文件存储"], [
        ["开发环境", "Vite dev server", "NestJS 本地服务", "本地 PostgreSQL / Docker", "本地 uploads"],
        ["测试环境", "静态构建产物", "Docker / PM2", "云 PostgreSQL", "OSS 测试 Bucket"],
        ["生产环境", "Nginx / OSS 静态站点", "阿里云 ECS / Docker", "阿里云 RDS PostgreSQL", "阿里云 OSS"],
    ], widths=[1.1, 1.5, 1.5, 1.5, 1.5])

    heading(doc, "10. 开发阶段规划", 1)
    add_table(doc, ["阶段", "目标", "交付物"], [
        ["阶段 1：前端原型", "跑通页面和主流程", "React 页面、静态题库 JSON、本地登录、假 AI 数据。"],
        ["阶段 2：后端与数据库", "接入真实用户和数据保存", "NestJS API、Prisma schema、PostgreSQL、历史记录。"],
        ["阶段 3：录音与 AI", "接入真实录音上传和 DeepSeek", "MediaRecorder、Upload API、AI 点评、生题、报告。"],
        ["阶段 4：正式化", "可上线使用", "对象存储、错误监控、日志、权限预留、题库管理。"],
    ], widths=[1.4, 2.0, 3.1])

    heading(doc, "11. 重要技术决策", 1)
    add_table(doc, ["决策点", "结论"], [
        ["倒计时", "只显示总倒计时，不显示单题倒计时。"],
        ["移动端", "暂不支持移动端模拟面试，优先桌面端。"],
        ["AI", "使用 DeepSeek API，由后端代理调用。"],
        ["账号", "暂不做手机号/微信登录，先做用户名 + 邮箱。"],
        ["存储", "开发期本地，正式期对象存储。"],
        ["会员", "V1.0 不做，数据库和模块预留。"],
    ], widths=[1.6, 4.9])

    doc.save(DOCX)
    return DOCX


if __name__ == "__main__":
    print(build())
