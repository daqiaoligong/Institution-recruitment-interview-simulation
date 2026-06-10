from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
DOCS.mkdir(exist_ok=True)
OUT = DOCS / "沪面冲鸭开发实施计划.docx"


def setup_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    styles = doc.styles
    for name in ["Normal", "List Bullet", "List Number"]:
        styles[name].font.name = "微软雅黑"
        styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        styles[name].font.size = Pt(10.5)


def heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.color.rgb = RGBColor(17, 24, 39)


def para(doc: Document, text: str = "") -> None:
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.size = Pt(10.5)


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            run.font.size = Pt(10.5)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.size = Pt(8.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell_text(t.rows[0].cells[i], h, True)
        shade(t.rows[0].cells[i], "F2F4F7")
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cell_text(cells[i], value)
    doc.add_paragraph()


def code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    r.font.size = Pt(8.5)


def build() -> Path:
    doc = Document()
    setup_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("沪面冲鸭开发实施计划")
    run.bold = True
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(24)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("上海事业单位面试模拟网站｜从前端 MVP 到后端、数据库、AI 的实施路线")
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(102, 112, 133)
    para(doc, f"版本：V1.0    日期：{date.today().isoformat()}")

    heading(doc, "1. 项目目标与技术栈")
    para(doc, "本计划用于指导沪面冲鸭从当前 PRD、架构设计和页面设计图进入实际开发。路线采用先前端可演示闭环，再真实录音，再接后端数据库，再接 DeepSeek AI，最后正式化上线。")
    bullets(doc, [
        "前端：React + Vite + TypeScript + Zustand。",
        "后端：NestJS + Prisma。",
        "数据库：PostgreSQL。",
        "文件存储：开发期本地 uploads，正式期阿里云 OSS。",
        "AI：DeepSeek API，由后端代理调用。",
        "录音：MediaRecorder，音频格式 WebM/Opus。",
        "转写：MVP 使用 Web Speech API，正式版预留服务端 ASR。",
        "包管理：pnpm monorepo，同时兼容 npm workspaces。",
    ])

    heading(doc, "2. 文件结构设计")
    code_block(doc, r"""E:\PG\SHshiyebianmianshimoni
├─ apps
│  ├─ web
│  │  ├─ public\assets
│  │  └─ src
│  │     ├─ app
│  │     ├─ pages
│  │     ├─ components
│  │     ├─ features
│  │     ├─ stores
│  │     ├─ services
│  │     ├─ data
│  │     ├─ types
│  │     ├─ utils
│  │     └─ styles
│  └─ api
│     ├─ src
│     ├─ prisma
│     └─ uploads\audio
├─ packages\shared
├─ docs
├─ assets
├─ design_outputs
├─ architecture_outputs
├─ scripts
└─ 历史代码""")

    heading(doc, "3. 分阶段实施计划")
    table(doc, ["阶段", "目标", "主要交付物"], [
        ["Phase 0", "文档、工程和数据准备", "实施计划 DOCX、docs 归档、monorepo 结构、环境模板、初始题库数据。"],
        ["Phase 1", "前端静态原型", "8 个核心路由、10 张设计图对应 UI、静态题库和自由组题。"],
        ["Phase 2", "前端面试流程闭环", "听题/看题流程、总倒计时、历史记录、复盘页 mock AI。"],
        ["Phase 3", "真实录音与前端转写", "MediaRecorder 分段录音、IndexedDB 音频保存、Web Speech 转写。"],
        ["Phase 4", "后端与数据库", "NestJS、Prisma、PostgreSQL、Auth/Question/Interview/Upload API。"],
        ["Phase 5", "DeepSeek AI 接入", "单题点评、答题思路、整场报告、岗位 AI 生题。"],
        ["Phase 6", "正式化与上线准备", "OSS 适配、鉴权、限流、日志、部署说明、会员预留。"],
    ])

    heading(doc, "4. 前端模块设计")
    bullets(doc, [
        "页面：Login、Home、InterviewSetup、InterviewSession、InterviewReview、QuestionBank、JobProfile、Profile。",
        "Stores：authStore、questionStore、jobProfileStore、interviewStore、historyStore。",
        "Services：apiClient、localStorageService、indexedDbService、recorderService、speechService、mockAiService。",
        "交互规则：只显示总倒计时；听题模式不显示题干；看题展开态不显示考官背景；收起态显示沉浸式背景。",
    ])

    heading(doc, "5. 后端模块设计")
    bullets(doc, [
        "AuthModule：用户名 + 邮箱登录，JWT 预留。",
        "QuestionModule：题库分类树、套题、专属题型。",
        "JobProfileModule：岗位信息保存与读取。",
        "InterviewModule：创建面试、结束面试、历史记录、复盘详情。",
        "AnswerModule：保存题目快照、转写文本、录音元数据。",
        "UploadModule：开发期本地录音上传，正式期 OSS 适配。",
        "AIModule：DeepSeek 调用、Prompt 模板、失败兜底。",
        "MembershipModule：会员功能预留，不进入 MVP。",
    ])

    heading(doc, "6. 数据库落地计划")
    table(doc, ["表名", "用途", "关键要求"], [
        ["users", "用户基础信息", "email 唯一。"],
        ["job_profiles", "岗位信息", "每用户默认一份当前岗位信息。"],
        ["questions", "标准题库题目", "支持题型、地区、日期筛选。"],
        ["question_sets", "套题", "保存真题套题信息。"],
        ["question_set_items", "套题题目关系", "按 sort_order 保证题目顺序。"],
        ["custom_questions", "我的专属题型", "单用户最多 10 道。"],
        ["interviews", "一次面试", "保存模式、总时长、状态。"],
        ["interview_answers", "逐题作答", "必须保存 question_content_snapshot。"],
        ["audio_files", "录音元数据", "只保存 URL 和元数据，文件不入库。"],
        ["ai_reviews", "单题 AI 点评", "保存评语、思路、优点、不足、建议。"],
        ["ai_reports", "整场 AI 报告", "保存总分、维度分、总评。"],
        ["memberships", "会员预留", "MVP 不启用支付。"],
    ])

    heading(doc, "7. API 落地计划")
    table(doc, ["方法", "路径", "说明"], [
        ["POST", "/api/auth/register", "注册用户"],
        ["POST", "/api/auth/login", "登录"],
        ["GET", "/api/auth/me", "当前用户"],
        ["GET", "/api/questions/tree", "题库分类树"],
        ["GET", "/api/questions", "题目查询"],
        ["GET", "/api/question-sets/:id", "套题详情"],
        ["POST", "/api/questions/custom", "新增专属题"],
        ["PUT", "/api/questions/custom/:id", "编辑专属题"],
        ["DELETE", "/api/questions/custom/:id", "删除专属题"],
        ["GET", "/api/job-profile", "获取岗位信息"],
        ["POST", "/api/job-profile", "保存岗位信息"],
        ["POST", "/api/interviews", "创建面试"],
        ["GET", "/api/interviews", "历史记录"],
        ["GET", "/api/interviews/:id", "复盘详情"],
        ["POST", "/api/interviews/:id/answers", "保存单题作答"],
        ["POST", "/api/interviews/:id/finish", "结束面试"],
        ["POST", "/api/uploads/audio", "上传录音"],
        ["POST", "/api/ai/review-answer", "单题 AI 点评"],
        ["POST", "/api/ai/generate-questions", "岗位 AI 生题"],
        ["POST", "/api/ai/generate-report", "整场报告"],
    ])

    heading(doc, "8. AI 与录音实现计划")
    bullets(doc, [
        "MediaRecorder 每题生成独立 WebM/Opus Blob。",
        "IndexedDB 保存本地音频 Blob，localStorage 保存元数据索引。",
        "Web Speech API 使用 zh-CN 做实时转写，不支持时允许手动编辑。",
        "DeepSeek Prompt 分为单题点评、答题思路、岗位生题、整场报告四类。",
        "DeepSeek 失败不影响录音、转写和历史记录，提供重试。",
    ])

    heading(doc, "9. 测试计划")
    bullets(doc, [
        "页面测试：登录拦截、导航、题库三栏、岗位信息、面试页视觉规则、复盘、个人中心。",
        "流程测试：听题 3 题、看题 3 题、总倒计时自动结束、下一题保存、重新开始确认。",
        "录音测试：授权成功/拒绝、不支持 MediaRecorder、每题分段、播放、空录音、转写失败。",
        "后端测试：注册登录、题库、岗位、创建面试、上传音频、保存作答、查询历史。",
        "AI 测试：DeepSeek 点评、报告、生题、超时、非 JSON、重试、失败兜底。",
    ])

    heading(doc, "10. 里程碑与验收")
    table(doc, ["里程碑", "验收标准"], [
        ["1 文档与工程骨架", "实施计划 Word 完成；monorepo 完成；前后端可启动。"],
        ["2 前端静态原型", "8 个核心页面完成；静态题库可浏览。"],
        ["3 前端本地闭环", "听题/看题流程跑通；本地历史可回看。"],
        ["4 录音转写闭环", "每题录音和转写可用；复盘页可播放。"],
        ["5 后端数据库闭环", "NestJS + Prisma + PostgreSQL 接通；录音上传可用。"],
        ["6 AI 闭环", "DeepSeek 单题点评、整场报告、岗位生题可用。"],
        ["7 测试部署", "测试环境部署；README 和环境变量完整；核心验收通过。"],
    ])

    heading(doc, "11. 风险与默认策略")
    bullets(doc, [
        "后端和存储尚未建设：先前端本地闭环，再接 NestJS/PostgreSQL。",
        "Web Speech API 兼容性有限：MVP 支持 Chrome/Edge，失败时允许手动编辑。",
        "音频不适合 localStorage：音频 Blob 用 IndexedDB，本地元数据用 localStorage。",
        "DeepSeek 可能失败或超时：AI 失败不阻塞复盘，提供重试。",
        "真题后续会更新：题库 seed 脚本和表结构支持增量导入。",
        "会员未来可能加入：先预留 memberships 表和模块，不做支付。",
        "旧代码不可靠：新版从 apps 开始，旧代码只参考。",
    ])

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
