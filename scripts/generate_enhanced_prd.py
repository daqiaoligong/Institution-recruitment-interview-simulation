from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "上海事业单位面试模拟网站PRD_完善版.docx"
DESIGN = ROOT / "design_outputs"
ASSETS = ROOT / "assets"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.size = Pt(9)
    r.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, True)
        set_cell_shading(hdr[i], "F2F4F7")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], str(val))
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.color.rgb = RGBColor(17, 24, 39)
    return p


def add_para(doc, text="", bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.name = "微软雅黑"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    for run in p.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.size = Pt(10.5)
    return p


def setup_styles(doc):
    styles = doc.styles
    styles["Normal"].font.name = "微软雅黑"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    styles["Normal"].font.size = Pt(10.5)
    for name in ["List Bullet", "List Number"]:
        styles[name].font.name = "微软雅黑"
        styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        styles[name].font.size = Pt(10.5)
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if (ASSETS / "logo.png").exists():
        r = p.add_run()
        r.add_picture(str(ASSETS / "logo.png"), width=Inches(1.25))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if (ASSETS / "name.png").exists():
        r = p.add_run()
        r.add_picture(str(ASSETS / "name.png"), width=Inches(3.4))
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("上海事业单位面试模拟网站 PRD")
    run.bold = True
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(24)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("完善版｜极简沉浸式面试模拟 + 题库 + AI 复盘")
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(102, 112, 133)
    meta = [
        "版本：V1.0 完善版",
        f"日期：{date.today().isoformat()}",
        "素材来源：PRD 草图、assets/logo.png、assets/name.png、assets/面试背景图.png、2025 上海事业单位面试真题 PDF、design_outputs 页面深化图",
    ]
    for m in meta:
        p = doc.add_paragraph(m)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()


def add_image(doc, path, caption):
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(path), width=Inches(6.65))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(102, 112, 133)


def build_doc():
    doc = Document()
    setup_styles(doc)
    add_cover(doc)

    add_heading(doc, "1. 文档说明", 1)
    add_para(doc, "本文档在原始 PRD 基础上，结合页面布局草图、已生成的 10 张深化版页面、品牌素材与 2025 年上海事业单位面试真题 PDF，对产品范围、功能需求、页面结构、数据需求、AI 能力、非功能性需求和验收标准进行补充。")
    add_para(doc, "当前文档定位为 MVP 到 V1.0 的产品需求说明，可供 UI 设计、前端开发、后端开发、AI 接口开发和测试验收共同使用。")

    add_heading(doc, "2. 产品概述", 1)
    add_para(doc, "产品名称：沪面冲鸭 - 上海事业单位面试模拟网站")
    add_para(doc, "产品定位：面向上海事业单位面试考生的在线模拟训练工具，提供真题题库、自由组题、听题/看题沉浸式模拟、录音转写、AI 点评、AI 答题思路和历史复盘。")
    add_para(doc, "核心价值：让考生在接近真实考场的环境中完成练习，并将每一次作答转化为可复盘、可改进、可沉淀的训练记录。")
    add_table(doc, ["目标用户", "核心诉求", "产品回应"], [
        ["备考上海事业单位面试的考生", "熟悉真实面试流程、积累真题、提升表达", "提供套题模拟、倒计时、听题/看题模式、录音复盘"],
        ["有明确报考岗位的考生", "围绕岗位匹配题和专业题专项训练", "岗位信息填报、AI 生题、岗位匹配评分"],
        ["需要复盘的高频练习用户", "保存每次练习结果并比较改进", "个人中心保存录音、转写文本、AI 评语和答题思路"],
    ], widths=[1.5, 2.3, 2.7])

    add_heading(doc, "3. 产品目标与范围", 1)
    add_heading(doc, "3.1 MVP 目标", 2)
    add_bullets(doc, [
        "完成登录/注册、本地或服务端用户身份识别。",
        "完成题库浏览、套题选择、自由组题和开始模拟。",
        "完成听题模式与看题模式两条面试流程。",
        "完成逐题录音、逐题转写、逐题保存。",
        "完成面试结束页，展示题目、录音、转写文本、AI 评语和 AI 答题思路。",
        "完成个人中心历史记录，支持回看每次模拟。",
    ])
    add_heading(doc, "3.2 V1.0 增强目标", 2)
    add_bullets(doc, [
        "岗位信息驱动 AI 生题，默认生成 10 道岗位匹配题/专业题。",
        "我的专属题型支持用户手动添加题目，单用户最多 10 道。",
        "AI 评分支持总评、逐题评语、答题思路、表达建议和岗位匹配分析。",
        "题库支持按上海市属、区属、年份、区域、单位、题型多维分类。",
        "报告支持导出或复制分享链接。",
    ])

    add_heading(doc, "4. 信息架构", 1)
    add_table(doc, ["一级页面", "入口", "页面职责"], [
        ["登录/注册页", "未登录访问、右上角个人入口", "完成用户注册、登录、建立用户身份"],
        ["开始面试首页", "默认首页/导航面试模拟", "展示品牌、进入面试配置"],
        ["面试配置页", "开始面试按钮后", "选择套题、每题时间、听题/看题模式"],
        ["模拟面试页", "配置页开始面试", "沉浸式面试、倒计时、录音、下一题、重新开始"],
        ["面试结束复盘页", "套题完成或倒计时结束", "逐题复盘、AI 评语、AI 答题思路"],
        ["题库页", "导航题库", "分类浏览题目、加入自由组题、开始自由模拟"],
        ["岗位信息页", "导航岗位信息", "填写岗位信息、AI 生题、加入我的专属题型"],
        ["个人中心页", "右上角个人图标", "查看历史面试记录和报告"],
    ], widths=[1.4, 1.8, 3.3])

    add_heading(doc, "5. 页面与交互需求", 1)
    add_heading(doc, "5.1 登录/注册页", 2)
    add_bullets(doc, [
        "用户输入用户名、邮箱即可注册或登录。MVP 阶段可采用邮箱唯一识别。",
        "未登录用户访问题库、模拟、岗位信息、个人中心时，应跳转登录页或弹出登录提示。",
        "登录成功后右上角显示个人图标，点击进入个人中心。",
        "需校验邮箱格式，用户名不可为空。",
    ])

    add_heading(doc, "5.2 开始面试首页", 2)
    add_bullets(doc, [
        "顶部显示 logo 与 name 品牌素材，导航包含面试模拟、题库、岗位信息、个人入口。",
        "页面中央展示品牌主视觉和“开始面试”按钮。",
        "风格要求：苹果风、极简、浅色背景、低饱和边框、少量品牌橙色强调。",
    ])

    add_heading(doc, "5.3 面试配置页", 2)
    add_bullets(doc, [
        "配置项包括：选择面试套题、选择每道题作答时间、选择听题/看题模式。",
        "套题来源包括题库套题、自由模拟套题、AI 生成题、我的专属题型。",
        "点击开始面试后进入全屏或近全屏的模拟面试页。",
        "若未选择套题，应提示用户先选择题目；若自由组题为空，不允许开始。",
    ])

    add_heading(doc, "5.4 模拟面试页 - 通用规则", 2)
    add_bullets(doc, [
        "面试页以 assets/面试背景图.png 为沉浸式考官背景。",
        "右上角显示小型总倒计时；总倒计时结束后自动结束本场面试并保存已完成内容。",
        "底部中央显示小型悬浮控件：录音状态、播放/暂停、暂停/继续、下一题、重新开始。",
        "控件不得使用大面积灰色衬底，不得遮挡考官背景；底部可使用轻微透明渐隐，面积应尽量小。",
        "录音状态需有持续可视化标记，如红点/橙点、脉冲或音量波动。",
        "下一题动作会结束当前题录音并保存当前题音频与转写文本。",
        "重新开始需二次确认，确认后清空本场临时录音与转写，回到第一题。",
    ])

    add_heading(doc, "5.5 听题模式", 2)
    add_bullets(doc, [
        "进入面试后直接显示考官背景，不展示题目正文。",
        "系统播放题目语音；读题期间总倒计时继续。",
        "读题结束后自动开始收音，或用户可手动控制开始/暂停。",
        "页面不显示“正在朗读题目”等中央提示卡，避免破坏沉浸感。",
        "用户回答结束后点击下一题，系统保存本题录音和转写文本，并播放下一题。",
    ])

    add_heading(doc, "5.6 看题模式", 2)
    add_bullets(doc, [
        "看题展开态：题目阅读层覆盖背景，不显示考官背景；但保留模式标签、倒计时和底部小控件。",
        "题目阅读层展示完整套题，支持点击最小化按钮收起题目。",
        "看题收起态：隐藏题目层，显示沉浸式考官背景，底部显示“显示题目”小按钮。",
        "看题模式从进入面试开始即开始录音；每次点击下一题时分段保存录音与转写。",
        "用户可在作答过程中随时显示/隐藏题目，但该操作不应影响录音计时。",
    ])

    add_heading(doc, "5.7 面试结束复盘页", 2)
    add_bullets(doc, [
        "左侧按题目分条展示：题号、题目、录音播放、转写文本摘要。",
        "点击某道题后，右侧展示该题 AI 评语和 AI 答题思路。",
        "页面应支持查看整场总评，包括总分、表达稳定度、岗位匹配度、结构化表达能力等。",
        "转写文本为空时，需要明确提示“未识别到有效作答”，AI 分析应基于空作答给出改进建议。",
    ])

    add_heading(doc, "5.8 题库页", 2)
    add_bullets(doc, [
        "采用三栏布局：左侧题库分类，中间题目内容，右侧自由组题。",
        "左侧一级目录：套题、题型分类题、岗位匹配题、我的专属题型。",
        "套题二级目录：上海市属事业单位、上海区属事业单位；上海区属事业单位下按区展示三级目录。",
        "题型分类题二级目录：综合分析、人际关系、应急应变、组织计划。",
        "点击套题后，中间展示该套题全部题目，每题右侧提供“加入模拟”。",
        "右侧自由组题支持添加、删除、编辑、粘贴自定义题目；编辑仅影响本次模拟，不回写原题库。",
        "自由组题区域右下角提供开始模拟按钮。",
    ])

    add_heading(doc, "5.9 岗位信息页", 2)
    add_bullets(doc, [
        "页面布局与题库页一致，左侧仍为题库分类，中间为岗位信息表单，右侧为 AI 生题展示区。",
        "岗位信息字段包括：岗位名称、岗位所在单位、岗位要求、岗位其他信息。",
        "保存后岗位信息用于 AI 生题、AI 评分、岗位匹配分析。",
        "点击 AI 生题后，默认生成 10 道可能考察的专业题/岗位匹配题。",
        "AI 生成题可编辑，用户确认后可加入“我的专属题型”。",
    ])

    add_heading(doc, "5.10 个人中心页", 2)
    add_bullets(doc, [
        "个人中心按列表展示历史面试记录，记录标题为面试时间。",
        "每条记录展示面试模式、题目数量、总时长、AI 报告状态。",
        "点击记录进入面试结束复盘页，查看当次题目、录音、转写文本和 AI 分析。",
        "历史记录支持按时间倒序排列。",
    ])

    add_heading(doc, "6. 题库与真题内容需求", 1)
    add_para(doc, "题库初始内容可参考 assets/【真题】2025年上海事业单位面试真题.pdf。该 PDF 共 9 页，包含 2025 年上海事业单位多场真题，覆盖市属单位、区属单位、储备人才、司法、环境、市容、职业技能鉴定、公积金、区级事业单位等场景。")
    add_table(doc, ["题库字段", "说明", "示例"], [
        ["题目 ID", "系统唯一标识", "q_20250529_ja_001"],
        ["题目标题", "用于列表展示", "2025年5月29日上海市静安区事业单位面试题"],
        ["题目正文", "完整面试题干", "人工智能是年轻的事业，也是年轻人的事业..."],
        ["题型", "综合分析/人际关系/应急应变/组织计划/岗位匹配/专业题", "综合分析"],
        ["来源类型", "真题/AI生成/用户自建", "真题"],
        ["地区", "上海市属或具体区", "静安区"],
        ["单位/岗位", "可为空", "司法第三中心"],
        ["年份/日期", "真题日期", "2025-05-29"],
        ["考场规则", "时间、题数、听题/看题、草稿纸等", "15分钟3道题，听题模式"],
        ["参考思路", "AI或人工整理", "用于复盘页展示"],
    ], widths=[1.2, 2.0, 3.2])
    add_heading(doc, "6.1 初始真题样例", 2)
    add_bullets(doc, [
        "2025年3月12日上海市自然资源部东海生态中心：15分钟4道题，读题听题模式，提供草稿纸和笔，无题本。",
        "2025年5月24日上海市市容环境质量监测中心：包含城市公厕建设、规划、居民参与等岗位专业题。",
        "2025年5月29日上海市静安区事业单位：包含人工智能、政务服务适老化、“两企三新”党建等上海热点题。",
        "2025年6月10日上海市杨浦区事业单位：包含守正创新、飞絮花粉争议、机构优化调整等题目。",
    ])

    add_heading(doc, "7. AI 能力需求", 1)
    add_table(doc, ["AI 功能", "输入", "输出", "要求"], [
        ["AI 转写后点评", "题目、转写文本、岗位信息", "逐题评语、优点、不足、改进建议", "语言具体，避免空泛套话"],
        ["AI 答题思路", "题目、题型、岗位信息", "结构化答题框架", "适合上海事业单位面试语境"],
        ["AI 总评", "整场题目与作答", "总分、维度分、综合建议", "支持空作答/短作答兜底"],
        ["AI 生题", "岗位名称、单位、要求、其他信息", "默认 10 道岗位匹配题/专业题", "题干可编辑，可加入题库"],
        ["岗位匹配分析", "岗位信息、作答内容", "人岗匹配度与提升建议", "突出岗位意识、服务意识、执行能力"],
    ], widths=[1.3, 1.6, 1.8, 1.8])
    add_bullets(doc, [
        "AI 输出需提示“仅供训练参考，不作为真实考试评分依据”。",
        "AI 失败时需展示本地兜底分析或提示稍后重试，不得阻断用户查看录音和转写文本。",
        "AI 接口调用需避免上传用户无关个人敏感信息。",
        "生成题加入题库前必须经用户确认。",
    ])

    add_heading(doc, "8. 录音与转写需求", 1)
    add_bullets(doc, [
        "需使用浏览器麦克风权限，首次使用时提示用户授权。",
        "每道题单独保存一段录音，音频记录需与题目 ID、用户 ID、本场面试 ID 关联。",
        "支持录音暂停/继续；暂停期间不计入有效录音片段，但总倒计时继续。",
        "转写可采用实时转写或结束后转写；MVP 可优先实时转写 + 保存最终文本。",
        "若浏览器不支持录音或转写，应提示用户更换 Chrome/Edge 或上传文字作答。",
        "录音文件应支持播放、暂停、进度拖动和时长展示。",
    ])

    add_heading(doc, "9. 数据与存储需求", 1)
    add_table(doc, ["数据对象", "关键字段", "说明"], [
        ["User", "id, username, email, createdAt", "用户基础信息"],
        ["JobProfile", "userId, title, unit, requirements, extraInfo, updatedAt", "岗位信息"],
        ["Question", "id, title, body, type, source, region, unit, date, rules", "题库题目"],
        ["MockSession", "id, userId, mode, totalTime, status, startedAt, endedAt", "一次模拟面试"],
        ["MockAnswer", "sessionId, questionId, audioUrl, transcript, duration, aiComment, aiThinking", "逐题作答记录"],
        ["CustomQuestion", "userId, body, type, createdAt", "我的专属题型，最多10道"],
    ], widths=[1.3, 3.0, 2.2])

    add_heading(doc, "10. 非功能性需求", 1)
    add_table(doc, ["类别", "需求"], [
        ["性能", "首页和核心页面首屏加载建议小于 2 秒；面试页切换题目响应小于 300ms；录音开始/停止反馈小于 500ms。"],
        ["稳定性", "倒计时、录音、下一题保存为核心链路，不应因 AI 接口失败而丢失用户作答。"],
        ["兼容性", "优先支持 Chrome、Edge 桌面端最新版；移动端可浏览题库和报告，模拟面试优先桌面端。"],
        ["隐私与安全", "录音、转写文本、岗位信息属于用户隐私；需明确授权、可删除、避免无必要外传。"],
        ["可用性", "面试页控件需小而清晰，避免遮挡背景；关键按钮需有明确状态反馈。"],
        ["可访问性", "文字对比度满足基本可读性；按钮点击区域不小于 44px；重要状态不只依赖颜色表达。"],
        ["可维护性", "题库、AI 提示词、评分维度、页面文案需配置化，避免硬编码在页面中。"],
        ["可扩展性", "后续可扩展更多地区、考试类型、付费题库、教师点评、报告导出。"],
        ["数据备份", "历史记录和录音需有备份策略；本地存储方案仅适合原型，不适合正式生产。"],
        ["合规", "真题内容需确认使用权；AI 生成内容需标注生成属性；不得承诺真实考试命中率。"],
    ], widths=[1.3, 5.2])

    add_heading(doc, "11. 页面视觉规范", 1)
    add_bullets(doc, [
        "整体风格：极简、清爽、苹果风，浅灰背景、白色面板、轻阴影、细边框。",
        "品牌素材：顶部使用 assets/logo.png 与 assets/name.png；启动页可展示 logo 吉祥物。",
        "品牌色：橙色用于强调、编号、录音状态；深色用于主按钮和导航选中态。",
        "面试沉浸页：背景图为核心视觉，不使用大面积衬底；控件仅作为悬浮工具出现。",
        "看题展开态：阅读题目时不显示考官背景，使用干净白底承载题目；保留小组件。",
    ])

    add_heading(doc, "12. 页面设计稿", 1)
    design_pages = [
        ("01_start_home.png", "图1 开始面试首页"),
        ("02_interview_setup.png", "图2 面试配置页"),
        ("03_listen_mode.png", "图3 听题模式沉浸页"),
        ("04_read_mode_expanded.png", "图4 看题模式 - 题目展开态"),
        ("05_read_mode_collapsed.png", "图5 看题模式 - 题目收起态"),
        ("06_interview_review.png", "图6 面试结束复盘页"),
        ("07_question_bank.png", "图7 题库页面"),
        ("08_job_info.png", "图8 岗位信息页面"),
        ("09_login.png", "图9 登录页面"),
        ("10_profile_history.png", "图10 个人中心页面"),
    ]
    for filename, caption in design_pages:
        add_image(doc, DESIGN / filename, caption)

    add_heading(doc, "13. 核心流程", 1)
    add_heading(doc, "13.1 听题模式流程", 2)
    add_numbered(doc, [
        "用户登录后进入面试配置页。",
        "选择套题、每题时间、听题模式，点击开始面试。",
        "系统进入沉浸式面试页，显示考官背景和小组件。",
        "系统播放第 1 题语音，总倒计时同步运行。",
        "读题结束后自动开始录音和转写。",
        "用户点击下一题，系统保存第 1 题录音和转写，进入第 2 题。",
        "全部题目完成或总倒计时结束后进入复盘页。",
        "系统生成 AI 评语、AI 答题思路和整场总评。",
    ])
    add_heading(doc, "13.2 看题模式流程", 2)
    add_numbered(doc, [
        "用户选择看题模式并开始面试。",
        "系统进入题目阅读层，显示完整套题，不显示考官背景。",
        "录音从进入面试起开始，倒计时同步运行。",
        "用户可点击最小化按钮收起题目，进入考官背景沉浸态。",
        "用户每答完一道题点击下一题，系统分段保存录音和转写。",
        "用户可随时点击显示题目回到阅读层。",
        "套题完成或总倒计时结束后进入复盘页。",
    ])

    add_heading(doc, "14. 验收标准", 1)
    add_table(doc, ["模块", "验收标准"], [
        ["登录注册", "用户可通过用户名和邮箱完成注册/登录；未登录访问核心功能会被拦截。"],
        ["面试配置", "可选择套题、时间、模式；配置缺失时有明确提示。"],
        ["听题模式", "不显示题目正文；可播放题目语音；读题后开始录音；下一题保存分段记录。"],
        ["看题模式", "展开态不显示背景；收起态显示考官背景；显示/隐藏题目不影响录音。"],
        ["录音转写", "每题生成独立录音和转写文本；录音状态有可视化反馈。"],
        ["复盘页", "逐题展示题目、录音、转写、AI评语和AI答题思路。"],
        ["题库页", "左侧分类、中间题目、右侧自由组题均可使用；可开始自由模拟。"],
        ["岗位信息", "可保存岗位信息；AI 生题默认 10 道；可编辑并加入我的专属题型。"],
        ["个人中心", "历史记录按时间倒序展示；可进入任意一次复盘。"],
        ["非功能", "核心链路不因 AI 失败而丢失录音/转写；主要浏览器可正常使用。"],
    ], widths=[1.5, 5.0])

    add_heading(doc, "15. 风险与待确认问题", 1)
    add_table(doc, ["问题", "建议默认方案", "需确认点"], [
        ["总倒计时与每题时间关系", "总倒计时由题数 x 每题时间计算；每题可不单独展示，只用于配置总时长", "是否需要同时显示单题倒计时？"],
        ["录音保存位置", "正式版使用服务端/对象存储；原型可用本地 Blob 临时保存", "是否已有后端和存储方案？"],
        ["AI 服务商", "接口抽象化，先支持一个模型服务，后续可替换", "使用 Gemini、OpenAI、国内模型还是自建？"],
        ["真题版权", "作为内部题库导入前需确认授权；页面显示来源信息", "PDF 真题是否允许公开商用？"],
        ["账号体系", "MVP 使用邮箱登录；正式版接入验证码或密码", "是否需要手机号/微信登录？"],
        ["移动端范围", "移动端优先支持浏览题库和报告，模拟面试优先桌面端", "是否必须支持手机面试模拟？"],
        ["付费与权限", "V1.0 暂不做付费，预留题库权限字段", "是否有会员/付费题库计划？"],
    ], widths=[1.8, 2.9, 1.8])

    add_heading(doc, "16. 版本规划建议", 1)
    add_table(doc, ["阶段", "目标", "范围"], [
        ["原型版", "验证页面和主流程", "10 张页面、静态题库、模拟录音/转写演示"],
        ["MVP", "可真实练习并保存复盘", "登录、题库、听题/看题、录音、转写、AI点评、历史记录"],
        ["V1.0", "形成完整训练闭环", "岗位AI生题、我的专属题型、报告导出、题库管理"],
        ["V1.5+", "商业化和扩展", "会员题库、教师点评、多地区多考试类型、学习数据分析"],
    ], widths=[1.2, 1.9, 3.4])

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_doc())
